from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

doc = Document()

# --- Styles ---
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

def add_heading_styled(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0, 51, 102)
    return h

def add_table_from_data(headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)
    for row_data in rows:
        row = table.add_row().cells
        for i, val in enumerate(row_data):
            row[i].text = str(val)
            for p in row[i].paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)
    doc.add_paragraph()

def add_bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        p.add_run(text)
    else:
        p.add_run(text)

def add_para(text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    return p

def add_note(text, label="NOTE"):
    p = doc.add_paragraph()
    run = p.add_run(f"[{label}] ")
    run.bold = True
    run.font.color.rgb = RGBColor(204, 102, 0)
    p.add_run(text)

# ============================================================
# TITLE PAGE
# ============================================================
doc.add_paragraph()
doc.add_paragraph()
title = doc.add_heading('CRMS – RESUME PARSING &\nINTELLIGENCE SYSTEM', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in title.runs:
    run.font.color.rgb = RGBColor(0, 51, 102)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sub.add_run('Final Architecture Document v2.0 (Revised & Improved)')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(80, 80, 80)

doc.add_paragraph()
info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
info.add_run('Date: 26 February 2026\n').font.size = Pt(11)
info.add_run('Status: Final Architecture Specification').font.size = Pt(11)

doc.add_page_break()

# ============================================================
# SYSTEM OBJECTIVE
# ============================================================
add_heading_styled('SYSTEM OBJECTIVE', 1)
objectives = [
    'Parse resume PDF into structured machine-readable format',
    'Standardize skill representation via hierarchical taxonomy',
    'Generate semantic embeddings (section-level + global)',
    'Evaluate candidate profile against job description',
    'Score candidate using calibrated evidence-based logic',
    'Rank applicants per job',
    'Suggest shortlist with explainable reasoning',
    'Allow HR override at every decision point',
    'Support future ML-based calibration via feedback loop',
    'Provide full audit trail & explainability dashboard',
]
for i, obj in enumerate(objectives, 1):
    add_bullet(f'{obj}')

doc.add_page_break()

# ============================================================
# LAYER 1
# ============================================================
add_heading_styled('LAYER 1 – Eligibility Filter (Deterministic Engine)', 1)
add_para('Type: Rule-based backend logic with versioned rule sets', bold=True)

add_heading_styled('Input', 2)
add_bullet('Student profile (branch, CGPA, batch, backlog count, gap years)')
add_bullet('Job eligibility criteria (set by HR per job posting)')

add_heading_styled('Output', 2)
add_table_from_data(
    ['Status', 'Description'],
    [
        ['ELIGIBLE', 'Meets all criteria'],
        ['BORDERLINE_ELIGIBLE', 'Within 5% margin of one criterion — flagged for HR review'],
        ['INELIGIBLE', 'Does not meet criteria — hard reject with reason logged'],
    ]
)

add_heading_styled('Eligibility Criteria (Expanded)', 2)
add_table_from_data(
    ['Criterion', 'Type', 'Example'],
    [
        ['Branch', 'Exact match / list', 'CSE, IT, ECE'],
        ['Minimum CGPA', 'Numeric threshold', '>= 7.0'],
        ['Batch year', 'Exact / range', '2025, 2026'],
        ['Active backlogs', 'Count threshold', '0 allowed'],
        ['Backlog history', 'Count threshold', '<= 2 historical'],
        ['Gap year policy', 'Boolean / count', '<= 1 year gap allowed'],
        ['Lateral entry', 'Boolean', 'Allowed / Not allowed'],
    ]
)

add_heading_styled('Rule Versioning', 2)
add_note('Every time HR modifies eligibility criteria, the system snapshots the rule set and associates it with the job posting. This ensures full auditability.', 'IMPORTANT')

add_heading_styled('Why Rule-Based', 2)
add_bullet('Campus hiring requires strict compliance — no probabilistic decisions')
add_bullet('Full auditability — every rejection has a documented reason')
add_bullet('No bias introduction — same rules applied to every candidate')
add_bullet('Legally defensible — can justify every decision with rule + data')

add_heading_styled('Why NOT ML', 2)
add_bullet('Non-deterministic — cannot guarantee same decision for same input')
add_bullet('Legally risky — cannot explain rejection to a candidate')
add_bullet('Hard to justify — "the model said so" is not acceptable')

doc.add_page_break()

# ============================================================
# LAYER 2
# ============================================================
add_heading_styled('LAYER 2 – Resume Parsing Layer (Revised)', 1)
add_para('Architecture: Hybrid Section-Based + NER Pipeline', bold=True)

add_heading_styled('Pipeline Flow', 2)
add_para('Resume PDF → Text Extraction (PyMuPDF + pdfplumber) → OCR Fallback (EasyOCR/Tesseract) → Section Classifier (Rule-Based) → Per-Section NER (Fine-tuned DistilBERT) → Structured JSON + Confidence Scores + Parse Status')

add_heading_styled('Technologies', 2)
add_table_from_data(
    ['Component', 'Technology', 'Purpose'],
    [
        ['PDF Text Extraction', 'PyMuPDF', 'Primary text extractor — fast, reliable'],
        ['Column/Table Handling', 'pdfplumber', 'Handles 2-column resumes, tables'],
        ['OCR (scanned PDFs)', 'EasyOCR (primary)', 'Better accuracy on styled documents'],
        ['OCR Fallback', 'Tesseract', 'Fallback for edge cases'],
        ['Text Cleaning', 'spaCy', 'Tokenization, sentence segmentation'],
        ['Section Detection', 'Rule-based classifier', 'Regex + keyword matching for headings'],
        ['Entity Extraction', 'Fine-tuned DistilBERT', 'Per-section NER with custom tags'],
    ]
)

add_heading_styled('Section Classifier (Pre-NER Step)', 2)
add_para('Before running NER, detect resume sections using heading-pattern matching for: Education, Skills, Projects, Experience, Certifications, Achievements.')
add_bullet('Provides structural context to NER')
add_bullet('Acts as fallback — even if NER fails, section data is still usable')
add_bullet('Improves extraction accuracy from ~40% to 80%+')

add_heading_styled('Fine-tuned DistilBERT NER Model', 2)
add_note('You MUST fine-tune DistilBERT on resume-specific labeled data. The base model does NOT know SKILL, PROJECT, CERTIFICATION entities.', 'CRITICAL')

add_heading_styled('Custom Entity Tags (BIO Schema)', 3)
add_table_from_data(
    ['Entity', 'Tag', 'Example'],
    [
        ['Skill', 'B-SKILL, I-SKILL', 'React.js, Machine Learning'],
        ['Project Name', 'B-PROJECT, I-PROJECT', 'E-Commerce Platform'],
        ['Certification', 'B-CERT, I-CERT', 'AWS Cloud Practitioner'],
        ['Organization', 'B-ORG, I-ORG', 'Google, TCS'],
        ['Degree', 'B-DEGREE, I-DEGREE', 'B.Tech CSE'],
        ['Date', 'B-DATE, I-DATE', 'June 2024'],
        ['Internship Role', 'B-ROLE, I-ROLE', 'Software Intern'],
        ['Technology', 'B-TECH, I-TECH', 'TensorFlow, PostgreSQL'],
    ]
)

add_heading_styled('Training Data Requirements', 3)
add_table_from_data(
    ['Item', 'Details'],
    [
        ['Required labeled resumes', '500–1,000 (minimum 500)'],
        ['Annotation tool', 'Label Studio (free) or Prodigy'],
        ['Training framework', 'HuggingFace Transformers'],
        ['Base model', 'distilbert-base-uncased'],
        ['Fine-tuning epochs', '5–10'],
        ['Target F1 score', '>= 0.85 per entity type'],
    ]
)

add_heading_styled('Why DistilBERT (Fine-tuned)?', 2)
add_table_from_data(
    ['Advantage', 'Detail'],
    [
        ['40% smaller than BERT', '66M vs 110M parameters'],
        ['Faster inference', '~2x faster than BERT-base'],
        ['CPU deployable', 'No GPU required'],
        ['Context-aware', 'Understands contextual meaning'],
        ['Open-source', 'No license cost'],
        ['No API dependency', 'Self-hosted, predictable cost'],
    ]
)

add_heading_styled('Why NOT Others?', 2)
add_table_from_data(
    ['Model', 'Reason Rejected'],
    [
        ['Regex only', 'Format fragile'],
        ['spaCy small NER', 'Limited contextual depth'],
        ['BERT-base', 'Heavy — 110M params, slower inference'],
        ['RoBERTa', 'Overkill for campus resumes'],
        ['LayoutLM', 'Unnecessary for semi-structured text'],
        ['GPT / LLM', 'Expensive, non-deterministic, API dependency'],
        ['DistilBERT (fine-tuned)', 'OPTIMAL cost-performance balance'],
    ]
)

add_heading_styled('Confidence Scoring', 2)
add_table_from_data(
    ['Confidence', 'Action'],
    [
        ['>= 0.85', 'Auto-accept'],
        ['0.60 – 0.84', 'Accept with flag for review'],
        ['< 0.60', 'Discard or flag for manual review'],
    ]
)

add_heading_styled('Parse Status Tracking', 2)
add_table_from_data(
    ['Status', 'Meaning'],
    [
        ['PARSED', 'All sections detected, NER successful'],
        ['PARTIAL_PARSE', 'Some sections missing or low confidence'],
        ['FAILED', 'Text extraction failed or NER returned empty'],
        ['OCR_FALLBACK', 'Primary extraction failed, OCR was used'],
    ]
)
add_note('Never silently fail. Every PARTIAL_PARSE and FAILED resume must surface on the HR dashboard.', 'WARNING')

doc.add_page_break()

# ============================================================
# LAYER 3
# ============================================================
add_heading_styled('LAYER 3 – Skill Intelligence Layer (Revised)', 1)
add_para('Technology: Hierarchical Skill Taxonomy stored in JSON / Database with admin UI', bold=True)

add_heading_styled('Purpose', 2)
add_bullet('Normalize skill names (synonyms, abbreviations, variations)')
add_bullet('Map related technologies into hierarchical categories')
add_bullet('Avoid duplication in scoring')
add_bullet('Enable partial credit matching via hierarchy')
add_bullet('Track skill proficiency levels when available')

add_heading_styled('Synonym & Abbreviation Mapping', 2)
add_table_from_data(
    ['Raw Text (from resume)', 'Normalized Skill'],
    [
        ['JS', 'JavaScript'], ['ML', 'Machine Learning'], ['DL', 'Deep Learning'],
        ['K8s', 'Kubernetes'], ['Postgres', 'PostgreSQL'], ['py', 'Python'],
        ['sklearn', 'scikit-learn'], ['React', 'React.js'], ['Node', 'Node.js'],
        ['Mongo', 'MongoDB'], ['tf', 'TensorFlow'], ['cpp', 'C++'],
    ]
)

add_heading_styled('Skill Proficiency Levels', 2)
add_table_from_data(
    ['Level', 'Keywords Detected'],
    [
        ['Expert', '"expert", "mastery", "advanced proficiency"'],
        ['Advanced', '"advanced", "proficient", "strong"'],
        ['Intermediate', '"intermediate", "working knowledge", "familiar"'],
        ['Beginner', '"beginner", "basic", "exposure", "learning"'],
        ['Unspecified', 'No indicator found (default)'],
    ]
)

add_heading_styled('Partial Credit via Hierarchy', 2)
add_table_from_data(
    ['JD Requirement', 'Candidate Has', 'Credit'],
    [
        ['React.js', 'React.js', '100% (exact match)'],
        ['JavaScript Framework', 'React.js', '90% (child match)'],
        ['JavaScript', 'React.js', '70% (parent skill)'],
        ['Frontend Development', 'React.js', '50% (category match)'],
        ['Web Development', 'React.js', '30% (broad category)'],
        ['Python', 'React.js', '0% (unrelated)'],
    ]
)

add_heading_styled('Why NOT ML Here?', 2)
add_bullet('Skill normalization is a deterministic ontology problem')
add_bullet('Business must control what counts as a valid skill')
add_bullet('Taxonomy updates must be instant and auditable')

doc.add_page_break()

# ============================================================
# LAYER 4
# ============================================================
add_heading_styled('LAYER 4 – Semantic Matching Layer (Revised)', 1)
add_para('Model Used: Sentence-BERT (all-MiniLM-L6-v2)', bold=True)

add_heading_styled('Section-Level Embedding Strategy', 2)
add_note('Do NOT embed the entire resume as one blob. Embed structured sections separately for dramatically better matching accuracy.', 'IMPORTANT')

add_para('Resume Embedding Strategy:')
add_bullet('skills_embedding = SBERT.encode(skills_text) → 384-dim vector')
add_bullet('projects_embedding = SBERT.encode(projects_text) → 384-dim vector')
add_bullet('experience_embedding = SBERT.encode(experience_text) → 384-dim vector')
add_bullet('overall_embedding = weighted_average(skills×0.35, projects×0.40, experience×0.25)')

add_para('JD Embedding Strategy:')
add_bullet('jd_skills_embedding = SBERT.encode(required_skills_text)')
add_bullet('jd_role_embedding = SBERT.encode(role_description_text)')
add_bullet('jd_overall_embedding = weighted_average(skills×0.40, role×0.60)')

add_heading_styled('Similarity Computation', 2)
add_bullet('skill_similarity = cosine(resume.skills_emb, jd.skills_emb)')
add_bullet('project_similarity = cosine(resume.projects_emb, jd.role_emb)')
add_bullet('experience_similarity = cosine(resume.experience_emb, jd.role_emb)')
add_bullet('global_similarity = cosine(resume.overall_emb, jd.overall_emb)')

add_heading_styled('JD Embedding Cache', 2)
add_bullet('JDs don\'t change → compute embedding once, store in Redis/DB')
add_bullet('Only compute resume embeddings at upload time')
add_bullet('Saves ~50% compute per matching operation')

add_heading_styled('Similarity Threshold Filtering', 2)
add_table_from_data(
    ['Global Cosine Similarity', 'Action'],
    [
        ['>= 0.60', 'Strong match — proceed to scoring'],
        ['0.30 – 0.59', 'Moderate match — score but flag'],
        ['< 0.30', 'Weak match — flag as "Likely Irrelevant"'],
    ]
)

add_heading_styled('Why Sentence-BERT?', 2)
add_table_from_data(
    ['Advantage', 'Detail'],
    [
        ['Designed for similarity', 'Purpose-built for semantic textual similarity'],
        ['Lightweight', '~22M parameters'],
        ['Fast inference', '<50ms per encoding on CPU'],
        ['CPU-friendly', 'No GPU required'],
        ['Fine-tunable', 'Can be fine-tuned later (Layer 6)'],
        ['No API cost', 'Self-hosted, open-source'],
        ['Deterministic', 'Same input → same output'],
    ]
)

add_heading_styled('Why NOT Alternatives?', 2)
add_table_from_data(
    ['Model', 'Reason Rejected'],
    [
        ['TF-IDF', 'Keyword-based — misses semantic relationships'],
        ['Word2Vec', 'No sentence-level context'],
        ['BERT embeddings (CLS)', 'Heavy, not optimized for similarity'],
        ['GPT embeddings', 'API cost, vendor dependency'],
        ['Cross-Encoder', 'Slow inference — doesn\'t scale'],
        ['Sentence-BERT', 'BEST similarity efficiency'],
    ]
)

doc.add_page_break()

# ============================================================
# LAYER 5
# ============================================================
add_heading_styled('LAYER 5 – Evidence-Based Scoring Engine (Revised)', 1)
add_para('Type: Deterministic 6-dimension weighted scoring with per-job-type configurable weights', bold=True)

add_heading_styled('Scoring Dimensions', 2)
add_table_from_data(
    ['#', 'Dimension', 'Default Weight', 'Source Layer'],
    [
        ['1', 'Skill Coverage', '20%', 'Layer 3'],
        ['2', 'Project Relevance & Complexity', '25%', 'Layer 2 + Layer 4'],
        ['3', 'Internship/Experience Relevance', '15%', 'Layer 2 + Layer 4'],
        ['4', 'Education Quality', '15%', 'Layer 1 + Layer 2'],
        ['5', 'Global Semantic Alignment', '10%', 'Layer 4'],
        ['6', 'Authenticity & Consistency', '15%', 'Cross-section analysis'],
    ]
)

add_heading_styled('Project Complexity Rubric', 3)
add_table_from_data(
    ['Signal', 'Points'],
    [
        ['Uses 3+ technologies', '+2'],
        ['Has deployment (live URL)', '+3'],
        ['Has GitHub / source link', '+2'],
        ['Team project', '+1'],
        ['Duration > 2 months', '+1'],
        ['ML/AI/Cloud component', '+2'],
        ['Real-world application', '+1'],
        ['Solo full-stack project', '+2'],
        ['Hackathon project', '+1'],
    ]
)

add_heading_styled('Education Quality Scoring', 3)
add_table_from_data(
    ['Signal', 'Points'],
    [
        ['CGPA >= 9.0', '+4'],
        ['CGPA 8.0–8.9', '+3'],
        ['CGPA 7.0–7.9', '+2'],
        ['Tier-1 institution (IIT/NIT/IIIT/BITS)', '+3'],
        ['Tier-2 institution', '+2'],
        ['Relevant major', '+2'],
        ['Relevant coursework mentioned', '+1'],
    ]
)

add_heading_styled('Authenticity & Consistency Checks', 3)
add_table_from_data(
    ['Check', 'Detection Method', 'Points'],
    [
        ['Date inconsistency', 'Overlapping dates', '-3'],
        ['Skill-project mismatch', 'Claims skill but no project uses it', '-2'],
        ['Buzzword stuffing', '>15 skills with no evidence', '-2'],
        ['Consistent timeline', 'Logical flow', '+2'],
        ['Evidence-backed claims', 'Skills appear in projects', '+3'],
        ['Verifiable links', 'GitHub, LinkedIn present', '+2'],
    ]
)

add_heading_styled('Per-Job-Type Weight Profiles', 2)
add_table_from_data(
    ['Role Type', 'Skill', 'Project', 'Intern', 'Edu', 'Semantic', 'Auth'],
    [
        ['Software Engineer', '20%', '25%', '15%', '15%', '10%', '15%'],
        ['Data Scientist / ML', '25%', '30%', '10%', '15%', '10%', '10%'],
        ['Web Developer', '25%', '30%', '10%', '10%', '10%', '15%'],
        ['Management Trainee', '10%', '15%', '25%', '20%', '15%', '15%'],
        ['Research Intern', '15%', '30%', '10%', '25%', '10%', '10%'],
    ]
)

doc.add_page_break()

# ============================================================
# LAYER 6
# ============================================================
add_heading_styled('LAYER 6 – Feedback & Calibration Layer (Revised)', 1)

add_heading_styled('Outcome Tracking', 2)
add_table_from_data(
    ['Stage', 'Data Captured'],
    [
        ['Applied', 'Resume uploaded, parsed, scored'],
        ['Shortlisted', 'Score-based shortlist + HR overrides'],
        ['Interviewed', 'Interview feedback (structured)'],
        ['Selected', 'Final offer extended'],
        ['Rejected', 'Stage of rejection + reason'],
    ]
)

add_heading_styled('Feedback Loop Process', 2)
add_bullet('Step 1: After hiring round, HR marks final outcomes')
add_bullet('Step 2: System runs correlation analysis (score vs selection)')
add_bullet('Step 3: Generate weight adjustment suggestions')
add_bullet('Step 4: HR reviews and approves/rejects adjustments')
add_bullet('Step 5: New weight profile versioned and stored')
add_bullet('Step 6: Next batch scored with new weights')

add_heading_styled('A/B Testing for Weights', 2)
add_bullet('Score candidates with BOTH old and new weights')
add_bullet('Compare rank changes and false positive/negative rates')
add_bullet('HR reviews comparison → Approves or rejects switch')

add_heading_styled('Supervised Ranking Transition Path', 2)
add_table_from_data(
    ['Phase', 'Requirement', 'Approach'],
    [
        ['Phase 1 (Current)', 'Launch', 'Deterministic scoring with fixed weights'],
        ['Phase 2', '~5,000 labeled outcomes', 'Logistic regression on outcomes'],
        ['Phase 3', '~20,000 outcomes', 'Fine-tune Sentence-BERT on domain data'],
        ['Phase 4', '~50,000 outcomes', 'LambdaMART / Learning-to-Rank model'],
    ]
)

doc.add_page_break()

# ============================================================
# RANKING & SHORTLISTING
# ============================================================
add_heading_styled('RANKING LOGIC', 1)
add_bullet('Calculate raw scores for all eligible candidates per job')
add_bullet('Compute z-scores: z = (score - mean) / std_dev')
add_bullet('Convert to percentile ranking')
add_bullet('Sort descending, assign ranks')
add_bullet('Store in applications table')

add_heading_styled('SHORTLISTING LOGIC', 1)
add_para('Recommended: Hybrid Model (Option C)', bold=True)
add_bullet('Shortlist if: percentile_rank >= dynamic_threshold OR rank <= top_X_count')
add_bullet('dynamic_threshold = max(70th percentile, minimum_quality_score)')
add_bullet('top_X_count = min(interview_slots, 20% of applicants)')

add_heading_styled('HR Override — Always Enabled', 2)
add_table_from_data(
    ['Override Type', 'Description'],
    [
        ['Force Include', 'HR adds candidate regardless of score'],
        ['Force Exclude', 'HR removes candidate with documented reason'],
        ['Adjust Threshold', 'HR raises/lowers threshold per job'],
        ['Override Score', 'HR adjusts dimension score with justification'],
    ]
)
add_note('Every HR override is logged with timestamp, HR name, and justification for audit.', 'IMPORTANT')

doc.add_page_break()

# ============================================================
# CROSS-CUTTING CONCERNS
# ============================================================
add_heading_styled('CROSS-CUTTING CONCERNS', 1)

add_heading_styled('1. Batch Processing Architecture', 2)
add_para('Upload → Job Queue (Celery + Redis) → Workers (Parse → NER → Embed → Score) → DB → Notify HR')
add_bullet('Processing Time per Resume: ~10-15 seconds (CPU)')
add_bullet('Throughput: ~300-400 resumes/hour on t3.large')

add_heading_styled('2. Error Handling & Graceful Degradation', 2)
add_table_from_data(
    ['Failure Point', 'Fallback'],
    [
        ['PyMuPDF fails', 'Try pdfplumber → OCR → Mark FAILED'],
        ['Section detection fails', 'Process entire text as one section'],
        ['NER returns empty', 'Use section-based heuristic extraction'],
        ['Embedding timeout', 'Retry → flag for review'],
        ['Scoring error', 'Log, exclude from ranking, notify HR'],
    ]
)

add_heading_styled('3. Resume De-duplication', 2)
add_bullet('Compute embedding of new resume and compare against existing')
add_bullet('If cosine_similarity > 0.95 → flag as near-duplicate')
add_bullet('Flag similar resumes from different students for plagiarism review')

add_heading_styled('4. Security & File Validation', 2)
add_table_from_data(
    ['Check', 'Action'],
    [
        ['File type', 'Only accept .pdf files'],
        ['File size', 'Max 5 MB per resume'],
        ['Malware scan', 'ClamAV scan before processing'],
        ['Text sanitization', 'Strip executable content'],
        ['PII handling', 'Encrypt sensitive data at rest'],
        ['Access control', 'Role-based (HR view, admin configure)'],
    ]
)

add_heading_styled('5. Explainability Dashboard', 2)
add_table_from_data(
    ['View', 'Content'],
    [
        ['Per-Candidate', 'Full score breakdown with explanations'],
        ['Per-Job Overview', 'Score distribution, shortlist recommendations'],
        ['Comparison', 'Side-by-side comparison of candidates'],
        ['Parsing Health', '% parsed successfully, common failures'],
        ['Taxonomy Stats', 'Most common skills, skill gaps'],
        ['Feedback Status', 'Weight calibration suggestions pending'],
    ]
)

doc.add_page_break()

# ============================================================
# TECHNOLOGY STACK
# ============================================================
add_heading_styled('TECHNOLOGY STACK', 1)
add_table_from_data(
    ['Component', 'Technology', 'Purpose'],
    [
        ['Backend API', 'FastAPI (Python)', 'REST API'],
        ['Job Queue', 'Celery + Redis', 'Async processing'],
        ['Database', 'PostgreSQL', 'Structured data'],
        ['Cache', 'Redis', 'Embedding & JD cache'],
        ['PDF Extraction', 'PyMuPDF + pdfplumber', 'Text extraction'],
        ['OCR', 'EasyOCR + Tesseract', 'Scanned documents'],
        ['NLP', 'spaCy', 'Text cleaning'],
        ['NER Model', 'Fine-tuned DistilBERT', 'Entity extraction'],
        ['Embedding Model', 'Sentence-BERT (MiniLM-L6)', 'Semantic similarity'],
        ['Frontend', 'React.js + Vite', 'HR Dashboard'],
        ['Monitoring', 'Prometheus + Grafana', 'System health'],
        ['File Storage', 'AWS S3 / GCP', 'Resume storage'],
        ['Deployment', 'Docker + Docker Compose', 'Containerization'],
        ['CI/CD', 'GitHub Actions', 'Automation'],
    ]
)

doc.add_page_break()

# ============================================================
# COST BREAKDOWN
# ============================================================
add_heading_styled('COST BREAKDOWN (Revised)', 1)

add_heading_styled('Monthly Infrastructure Cost', 2)
add_table_from_data(
    ['Component', 'Service', 'Monthly Cost (INR)'],
    [
        ['Compute Server', 'AWS EC2 t3.large', '₹6,000'],
        ['Database', 'PostgreSQL RDS', '₹2,000'],
        ['Cache', 'Redis', '₹1,200'],
        ['File Storage', 'S3', '₹500'],
        ['Monitoring', 'CloudWatch / Prometheus', '₹500'],
        ['Backup & DR', 'Automated backups', '₹500'],
        ['SSL / Domain', 'Certificate + DNS', '₹200'],
        ['Total (without LLM)', '', '₹10,900/month'],
        ['Optional: LLM', 'OpenAI API', '₹1,200'],
        ['Total (with LLM)', '', '₹12,100/month'],
    ]
)

add_heading_styled('Cost Per Resume', 2)
add_table_from_data(
    ['Scenario', 'Monthly Cost', 'Resumes/Month', 'Cost/Resume'],
    [
        ['Without LLM', '₹10,900', '3,000', '₹3.63'],
        ['With LLM', '₹12,100', '3,000', '₹4.03'],
        ['High volume', '₹10,900', '6,000', '₹1.82'],
    ]
)

add_heading_styled('Model Compute Cost', 2)
add_para('DistilBERT + Sentence-BERT: Both open-source (Apache 2.0). No API cost. No per-call charge. Model cost = ₹0 (included in server cost).', bold=True)

add_heading_styled('Implementation Timeline', 2)
add_table_from_data(
    ['Phase', 'Effort', 'Description'],
    [
        ['Phase 1: Core Pipeline', '2–3 weeks', 'Layers 1–4 implementation'],
        ['Phase 2: Scoring + Dashboard', '2–3 weeks', 'Layer 5 + Explainability UI'],
        ['Phase 3: Feedback Loop', '1–2 weeks', 'Layer 6 + A/B testing'],
        ['Phase 4: NER Fine-tuning', '2–3 weeks', 'Data labeling + model training'],
        ['Total', '7–11 weeks', ''],
    ]
)

doc.add_page_break()

# ============================================================
# PRIORITY MATRIX
# ============================================================
add_heading_styled('IMPLEMENTATION PRIORITY MATRIX', 1)
add_table_from_data(
    ['Priority', 'Task', 'Impact', 'Timeline'],
    [
        ['P0', 'Fine-tune DistilBERT on resume NER', 'Critical', 'Weeks 1–3'],
        ['P0', 'Section-based parsing pipeline', 'Critical', 'Week 1–2'],
        ['P0', 'Core API + Eligibility Filter', 'Critical', 'Week 1'],
        ['P1', 'Section-level embedding strategy', 'High', 'Week 3'],
        ['P1', 'Scoring engine with rubrics', 'High', 'Week 3–4'],
        ['P1', 'Explainability dashboard', 'High', 'Week 4–5'],
        ['P1', 'Job queue (Celery + Redis)', 'High', 'Week 2'],
        ['P2', 'Hierarchical skill taxonomy', 'Medium', 'Week 2–3'],
        ['P2', 'Per-job-type weight profiles', 'Medium', 'Week 4'],
        ['P2', 'JD embedding cache', 'Medium', 'Week 3'],
        ['P2', 'Resume de-duplication', 'Medium', 'Week 5'],
        ['P3', 'Taxonomy admin UI', 'Nice-to-have', 'Week 6–7'],
        ['P3', 'A/B testing for weights', 'Nice-to-have', 'Week 7'],
        ['P3', 'Parsing accuracy dashboard', 'Nice-to-have', 'Week 6'],
    ]
)

doc.add_page_break()

# ============================================================
# FINAL SUMMARY
# ============================================================
add_heading_styled('FINAL SUMMARY', 1)
add_table_from_data(
    ['Aspect', 'Detail'],
    [
        ['Architecture', '6-layer with cross-cutting concerns'],
        ['ML Models', 'DistilBERT (fine-tuned NER) + Sentence-BERT'],
        ['API Cost', '₹0 — all models self-hosted'],
        ['Monthly Infra', '₹10,900 (no LLM) / ₹12,100 (with LLM)'],
        ['Cost per Resume', '~₹3.63'],
        ['Processing Speed', '~10-15 sec/resume on CPU'],
        ['Throughput', '~300-400 resumes/hour'],
        ['Scalability', 'Horizontal scaling via Celery workers'],
        ['Explainability', 'Full scoring breakdown per candidate'],
        ['Auditability', 'Rule versioning, score logs, HR override logs'],
        ['Future-proof', 'Feedback loop → supervised ML transition'],
        ['Development Time', '7–11 weeks'],
    ]
)

add_para('')
add_note('Start with P0 tasks: Fine-tune DistilBERT, build section-based parser, create core API. Everything else builds on this foundation.', 'TIP')

add_para('')
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Document Version: 2.0 | Last Updated: 26 February 2026 | Status: Final Architecture Specification')
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(128, 128, 128)

# ============================================================
# SAVE
# ============================================================
output_path = r'c:\Dev\CRMS_Resume\CRMS_Final_Architecture_v2.docx'
doc.save(output_path)
print(f'Document saved to: {output_path}')
