from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

def heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for r in h.runs:
        r.font.color.rgb = RGBColor(0, 51, 102)
    return h

def table(headers, rows):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = 'Light Grid Accent 1'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        t.rows[0].cells[i].text = h
        for p in t.rows[0].cells[i].paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(9)
    for rd in rows:
        row = t.add_row().cells
        for i, v in enumerate(rd):
            row[i].text = str(v)
            for p in row[i].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)
    doc.add_paragraph()

def bullet(text):
    doc.add_paragraph(text, style='List Bullet')

def para(text, bold=False, size=11, color=None):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    r.font.size = Pt(size)
    if color: r.font.color.rgb = RGBColor(*color)
    return p

def note(text, label="NOTE"):
    p = doc.add_paragraph()
    r = p.add_run(f"⚠ {label}: ")
    r.bold = True
    r.font.color.rgb = RGBColor(204, 80, 0)
    p.add_run(text)

# ==================== TITLE PAGE ====================
doc.add_paragraph()
doc.add_paragraph()
t = doc.add_heading('CRMS – RESUME PARSING &\nINTELLIGENCE SYSTEM', level=0)
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
for r in t.runs: r.font.color.rgb = RGBColor(0, 51, 102)

s = doc.add_paragraph()
s.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = s.add_run('Resource Access & Procurement Checklist')
r.font.size = Pt(16)
r.font.color.rgb = RGBColor(80, 80, 80)

doc.add_paragraph()
s2 = doc.add_paragraph()
s2.alignment = WD_ALIGN_PARAGRAPH.CENTER
s2.add_run('Prepared for: Project Manager\n').font.size = Pt(11)
s2.add_run('Date: 26 February 2026\n').font.size = Pt(11)
s2.add_run('Status: Pre-Development Resource Request\n').font.size = Pt(11)
s2.add_run('Project Timeline: 7–11 Weeks').font.size = Pt(11)

doc.add_page_break()

# ==================== TABLE OF CONTENTS ====================
heading('TABLE OF CONTENTS', 1)
toc = [
    '1. Cloud Infrastructure (AWS / GCP / Azure)',
    '2. Development Tools & Software Licenses',
    '3. AI / ML Resources & Models',
    '4. Data Requirements (MOST CRITICAL)',
    '5. Database & Storage Services',
    '6. Domain, Networking & Security',
    '7. Third-Party API Services',
    '8. Python Libraries & Open-Source Packages',
    '9. Hardware Requirements (Development Machines)',
    '10. Team Access & Stakeholders',
    '11. Compliance & Legal Approvals',
    '12. Cost Summary (Monthly + One-Time)',
    '13. Priority-Wise Summary for PM',
]
for item in toc:
    bullet(item)
doc.add_page_break()

# ==================== 1. CLOUD INFRASTRUCTURE ====================
heading('1. CLOUD INFRASTRUCTURE (AWS / GCP / Azure)', 1)
para('All services below are required for hosting the CRMS system in production.', bold=True)

heading('1.1 Compute Servers', 2)
table(
    ['#', 'Resource', 'Specification', 'Purpose', 'Est. Monthly Cost', 'Priority'],
    [
        ['1', 'Primary Server (API + ML)', 'EC2 t3.large — 2 vCPU, 8GB RAM, 50GB SSD', 'Run FastAPI backend, DistilBERT, Sentence-BERT, Celery workers', '₹6,000', 'P0'],
        ['2', 'GPU Instance (One-Time)', 'EC2 g4dn.xlarge — 1 GPU, 16GB VRAM', 'Fine-tune DistilBERT NER model (4–8 hours only)', '₹500 (one-time)', 'P0'],
        ['3', 'Staging Server (Optional)', 'EC2 t3.medium — 2 vCPU, 4GB RAM', 'Testing/staging environment', '₹3,500', 'P3'],
    ]
)

heading('1.2 Storage Services', 2)
table(
    ['#', 'Resource', 'Specification', 'Purpose', 'Est. Monthly Cost', 'Priority'],
    [
        ['1', 'Resume Storage Bucket', 'S3 Standard — up to 10GB initial', 'Store uploaded resume PDFs', '₹200', 'P0'],
        ['2', 'Model Storage Bucket', 'S3 Standard — ~2GB', 'Store fine-tuned DistilBERT & SBERT models', '₹100', 'P0'],
        ['3', 'Backup Storage Bucket', 'S3 Infrequent Access — ~5GB', 'Database & config backups', '₹100', 'P2'],
        ['4', 'Log Storage', 'S3 / CloudWatch Logs', 'Application & error logs', '₹100', 'P2'],
    ]
)

heading('1.3 Managed Services', 2)
table(
    ['#', 'Resource', 'Specification', 'Purpose', 'Est. Monthly Cost', 'Priority'],
    [
        ['1', 'PostgreSQL Database', 'RDS db.t3.micro (2 vCPU, 1GB RAM, 20GB)', 'Primary application database', '₹2,000', 'P0'],
        ['2', 'Redis Cache', 'ElastiCache cache.t3.micro (1 vCPU, 0.5GB)', 'Celery job queue + embedding cache', '₹1,200', 'P1'],
        ['3', 'CloudWatch Monitoring', 'Basic monitoring + alarms', 'Server health & alerts', '₹500', 'P2'],
    ]
)

heading('1.4 IAM & Security', 2)
table(
    ['#', 'Resource', 'What to Request', 'Priority'],
    [
        ['1', 'IAM Admin User', 'Full access account for infra setup', 'P0'],
        ['2', 'IAM Developer Users', 'Scoped access for each developer (S3, EC2, RDS)', 'P0'],
        ['3', 'IAM Service Account', 'Programmatic access keys for backend API', 'P0'],
        ['4', 'Security Groups', 'Firewall rules: 443 (HTTPS), 5432 (Postgres), 6379 (Redis)', 'P1'],
        ['5', 'KMS Encryption Key', 'For encrypting PII data at rest', 'P2'],
    ]
)

doc.add_page_break()

# ==================== 2. DEVELOPMENT TOOLS ====================
heading('2. DEVELOPMENT TOOLS & SOFTWARE LICENSES', 1)

heading('2.1 Version Control & CI/CD', 2)
table(
    ['#', 'Tool', 'License', 'Cost', 'Purpose', 'Priority'],
    [
        ['1', 'GitHub / GitLab Private Repo', 'Free (public) / Team plan', 'Free – ₹350/user/month', 'Source code management', 'P0'],
        ['2', 'GitHub Actions / GitLab CI', 'Included in repo', 'Free (2000 min/month)', 'CI/CD pipelines', 'P2'],
        ['3', 'Git (Desktop Client)', 'Free', '₹0', 'Version control', 'P0'],
    ]
)

heading('2.2 IDE & Editors', 2)
table(
    ['#', 'Tool', 'License', 'Cost', 'Priority'],
    [
        ['1', 'VS Code', 'Free (open-source)', '₹0', 'P0'],
        ['2', 'PyCharm Professional (Optional)', 'Paid license', '₹700/month per user', 'P3'],
        ['3', 'Cursor / Copilot (Optional)', 'AI coding assistant', '₹800–₹1,600/month', 'P3'],
    ]
)

heading('2.3 Containerization & Deployment', 2)
table(
    ['#', 'Tool', 'License', 'Cost', 'Purpose', 'Priority'],
    [
        ['1', 'Docker Desktop', 'Free (small teams)', '₹0', 'Build & run containers', 'P1'],
        ['2', 'Docker Hub / AWS ECR', 'Free tier / Pay-per-use', '₹0 – ₹500', 'Store Docker images', 'P1'],
        ['3', 'Docker Compose', 'Free (included with Docker)', '₹0', 'Multi-container orchestration', 'P1'],
    ]
)

heading('2.4 API Testing & Documentation', 2)
table(
    ['#', 'Tool', 'License', 'Cost', 'Priority'],
    [
        ['1', 'Postman', 'Free tier', '₹0', 'P2'],
        ['2', 'Swagger / OpenAPI', 'Free (auto-generated by FastAPI)', '₹0', 'P0'],
    ]
)

heading('2.5 Project Management', 2)
table(
    ['#', 'Tool', 'License', 'Cost', 'Priority'],
    [
        ['1', 'Jira / Linear / Trello', 'Free tier or team plan', '₹0 – ₹600/user/month', 'P1'],
        ['2', 'Slack / Microsoft Teams', 'Company provided', '₹0', 'P1'],
        ['3', 'Confluence / Notion (Docs)', 'Free tier', '₹0', 'P2'],
    ]
)

doc.add_page_break()

# ==================== 3. AI/ML RESOURCES ====================
heading('3. AI / ML RESOURCES & MODELS', 1)

heading('3.1 Pre-trained Models (FREE — Open Source)', 2)
table(
    ['#', 'Model', 'Source', 'License', 'Size', 'Purpose', 'Cost'],
    [
        ['1', 'distilbert-base-uncased', 'HuggingFace', 'Apache 2.0', '~260MB', 'Base model for NER fine-tuning', '₹0'],
        ['2', 'all-MiniLM-L6-v2', 'HuggingFace', 'Apache 2.0', '~80MB', 'Sentence embedding (Semantic matching)', '₹0'],
        ['3', 'spaCy en_core_web_sm', 'spaCy', 'MIT', '~12MB', 'Text cleaning & tokenization', '₹0'],
        ['4', 'spaCy en_core_web_md (Optional)', 'spaCy', 'MIT', '~43MB', 'Better NLP pipeline', '₹0'],
    ]
)
para('Total Model Cost: ₹0 (all open-source, self-hosted)', bold=True)

heading('3.2 Model Training Resources', 2)
table(
    ['#', 'Resource', 'Option A', 'Option B', 'Est. Cost', 'Priority'],
    [
        ['1', 'GPU for Fine-tuning', 'Google Colab Pro', 'AWS g4dn.xlarge (4–8 hrs)', '₹1,000 – ₹1,500 (one-time)', 'P0'],
        ['2', 'Label Studio (Annotation)', 'Self-hosted (free)', 'Label Studio Cloud', '₹0 – ₹3,000/month', 'P0'],
        ['3', 'HuggingFace Account', 'Free account', '—', '₹0', 'P0'],
        ['4', 'Weights & Biases (Experiment Tracking)', 'Free tier', '—', '₹0', 'P3'],
    ]
)

heading('3.3 OCR Resources', 2)
table(
    ['#', 'Tool', 'License', 'Cost', 'Purpose', 'Priority'],
    [
        ['1', 'EasyOCR', 'Apache 2.0', '₹0', 'Primary OCR for scanned resumes', 'P1'],
        ['2', 'Tesseract OCR', 'Apache 2.0', '₹0', 'Fallback OCR engine', 'P1'],
        ['3', 'Tesseract Language Data (eng)', 'Free', '₹0', 'English language pack', 'P1'],
    ]
)

doc.add_page_break()

# ==================== 4. DATA REQUIREMENTS ====================
heading('4. DATA REQUIREMENTS (MOST CRITICAL)', 1)
note('This section contains the single most important resource for the project. Without sample resumes, the NER model cannot be trained and the core parsing layer will not work.', 'CRITICAL')

heading('4.1 Training Data', 2)
table(
    ['#', 'Data Required', 'Quantity', 'Source', 'Format', 'Priority'],
    [
        ['1', 'Student Resumes (Anonymized)', '500 – 1,000', 'Placement Cell / Past Drives', 'PDF files', 'P0 ⭐'],
        ['2', 'Job Descriptions', '20 – 50', 'HR / Past Campus Hiring', 'Text / PDF / Word', 'P0 ⭐'],
        ['3', 'Student Profile Data', 'All students in system', 'Placement Cell Database', 'CSV / Excel / DB export', 'P0'],
        ['4', 'Historical Hiring Outcomes', 'As much as available', 'HR Records', 'CSV / Excel', 'P3'],
    ]
)

heading('4.2 Resume Data Specifications', 2)
bullet('Resumes should cover diverse formats: single-column, two-column, with tables, with images')
bullet('Mix of branches: CSE, IT, ECE, Mechanical, etc.')
bullet('Must be anonymized: replace real names with dummy names if required')
bullet('Include both strong and weak resumes for balanced training')
bullet('Variety of resume lengths: 1-page and 2-page resumes')

heading('4.3 Student Profile Data Fields Needed', 2)
table(
    ['#', 'Field', 'Type', 'Example', 'Required'],
    [
        ['1', 'Student ID', 'String', 'STU-2025-001', 'Yes'],
        ['2', 'Full Name', 'String', 'Rahul Sharma', 'Yes'],
        ['3', 'Email', 'String', 'rahul@email.com', 'Yes'],
        ['4', 'Branch', 'String', 'CSE / IT / ECE', 'Yes'],
        ['5', 'CGPA / Percentage', 'Decimal', '8.2 / 82%', 'Yes'],
        ['6', 'Batch Year', 'Integer', '2025', 'Yes'],
        ['7', 'Active Backlogs', 'Integer', '0', 'Yes'],
        ['8', 'Backlog History', 'Integer', '2', 'Recommended'],
        ['9', 'Gap Years', 'Integer', '0', 'Recommended'],
        ['10', 'Phone Number', 'String', '+91-9876543210', 'Optional'],
    ]
)

heading('4.4 Job Description Data Fields Needed', 2)
table(
    ['#', 'Field', 'Type', 'Example'],
    [
        ['1', 'Job Title', 'String', 'Software Development Engineer'],
        ['2', 'Company Name', 'String', 'Infosys'],
        ['3', 'Job Description', 'Text', 'Build scalable web applications...'],
        ['4', 'Required Skills', 'List', 'Python, React, SQL, Docker'],
        ['5', 'Eligible Branches', 'List', 'CSE, IT, ECE'],
        ['6', 'Minimum CGPA', 'Decimal', '7.0'],
        ['7', 'Batch Year', 'Integer', '2025'],
        ['8', 'CTC / Stipend', 'String', '₹12 LPA / ₹25,000/month'],
        ['9', 'Role Type', 'String', 'Full-time / Intern'],
    ]
)

doc.add_page_break()

# ==================== 5. DATABASE & STORAGE ====================
heading('5. DATABASE & STORAGE SERVICES', 1)
table(
    ['#', 'Service', 'Technology', 'Specification', 'Monthly Cost', 'Priority'],
    [
        ['1', 'Primary Database', 'PostgreSQL 15+', 'RDS db.t3.micro or self-hosted', '₹0 – ₹2,000', 'P0'],
        ['2', 'Cache & Queue', 'Redis 7+', 'ElastiCache or self-hosted', '₹0 – ₹1,200', 'P1'],
        ['3', 'File Storage', 'AWS S3', '3 buckets (resumes, models, backups)', '₹500', 'P0'],
        ['4', 'Vector Storage (Future)', 'pgvector extension', 'Installed in PostgreSQL', '₹0 (extension)', 'P3'],
    ]
)
para('Note: If self-hosting PostgreSQL and Redis on the same EC2, database cost = ₹0 additional.', bold=True)

doc.add_page_break()

# ==================== 6. DOMAIN & NETWORKING ====================
heading('6. DOMAIN, NETWORKING & SECURITY', 1)
table(
    ['#', 'Resource', 'Specification', 'Cost', 'Priority'],
    [
        ['1', 'Domain Name', 'e.g., crms.yourcompany.com', '₹500 – ₹1,500/year', 'P1'],
        ['2', 'SSL Certificate', 'Let\'s Encrypt (free) or commercial', '₹0 – ₹5,000/year', 'P1'],
        ['3', 'Elastic IP', 'AWS Elastic IP (static)', '₹0 (if attached to instance)', 'P1'],
        ['4', 'VPN Access', 'Company VPN for dev access', '₹0 (company provided)', 'P2'],
        ['5', 'CDN (Optional)', 'CloudFront for dashboard', '₹300 – ₹800/month', 'P3'],
        ['6', 'DDoS Protection', 'AWS Shield Basic', '₹0 (included)', 'P2'],
        ['7', 'WAF (Optional)', 'AWS WAF for API protection', '₹500/month', 'P3'],
    ]
)

doc.add_page_break()

# ==================== 7. THIRD-PARTY APIS ====================
heading('7. THIRD-PARTY API SERVICES', 1)
table(
    ['#', 'Service', 'Provider', 'Purpose', 'Monthly Cost', 'Priority'],
    [
        ['1', 'LLM API (Optional)', 'OpenAI GPT-4o-mini', 'AI candidate summaries & explanations', '₹1,000 – ₹2,000', 'P3'],
        ['2', 'Email Service', 'AWS SES / SendGrid', 'Send notifications to HR', '₹0 – ₹500', 'P2'],
        ['3', 'Error Tracking', 'Sentry (free tier)', 'Application error monitoring', '₹0', 'P2'],
        ['4', 'Analytics (Optional)', 'Mixpanel / PostHog (free)', 'Dashboard usage analytics', '₹0', 'P3'],
        ['5', 'ClamAV', 'Open-source', 'Malware scanning for uploaded PDFs', '₹0', 'P2'],
    ]
)

doc.add_page_break()

# ==================== 8. PYTHON LIBRARIES ====================
heading('8. PYTHON LIBRARIES & OPEN-SOURCE PACKAGES', 1)
para('All packages below are free and open-source. Listed here for completeness.', bold=True)

heading('8.1 Core Backend', 2)
table(
    ['Package', 'Version', 'License', 'Purpose'],
    [
        ['FastAPI', 'Latest', 'MIT', 'REST API framework'],
        ['Uvicorn', 'Latest', 'BSD', 'ASGI server'],
        ['Celery', 'Latest', 'BSD', 'Task queue for async processing'],
        ['Redis (py)', 'Latest', 'MIT', 'Redis client'],
        ['SQLAlchemy', 'Latest', 'MIT', 'ORM for PostgreSQL'],
        ['Alembic', 'Latest', 'MIT', 'Database migrations'],
        ['psycopg2', 'Latest', 'LGPL', 'PostgreSQL driver'],
        ['Pydantic', 'Latest', 'MIT', 'Data validation'],
        ['python-multipart', 'Latest', 'Apache 2.0', 'File upload handling'],
    ]
)

heading('8.2 PDF & OCR', 2)
table(
    ['Package', 'Version', 'License', 'Purpose'],
    [
        ['PyMuPDF (fitz)', 'Latest', 'AGPL / Commercial', 'PDF text extraction'],
        ['pdfplumber', 'Latest', 'MIT', 'Column/table extraction from PDFs'],
        ['EasyOCR', 'Latest', 'Apache 2.0', 'OCR for scanned documents'],
        ['pytesseract', 'Latest', 'Apache 2.0', 'Tesseract OCR wrapper'],
        ['Pillow', 'Latest', 'HPND', 'Image processing for OCR'],
    ]
)
note('PyMuPDF uses AGPL license. For commercial use, you may need a commercial license (~$599 one-time). Verify with legal team.', 'LICENSE CHECK')

heading('8.3 AI / ML', 2)
table(
    ['Package', 'Version', 'License', 'Purpose'],
    [
        ['transformers', 'Latest', 'Apache 2.0', 'DistilBERT model loading & inference'],
        ['torch (PyTorch)', 'Latest', 'BSD', 'Deep learning framework'],
        ['sentence-transformers', 'Latest', 'Apache 2.0', 'Sentence-BERT embeddings'],
        ['spaCy', 'Latest', 'MIT', 'NLP text processing'],
        ['scikit-learn', 'Latest', 'BSD', 'Cosine similarity & metrics'],
        ['numpy', 'Latest', 'BSD', 'Numerical computing'],
        ['pandas', 'Latest', 'BSD', 'Data manipulation'],
    ]
)

heading('8.4 Frontend (React Dashboard)', 2)
table(
    ['Package', 'Version', 'License', 'Purpose'],
    [
        ['React', '18+', 'MIT', 'UI framework'],
        ['Vite', 'Latest', 'MIT', 'Build tool'],
        ['Axios', 'Latest', 'MIT', 'HTTP client'],
        ['Recharts / Chart.js', 'Latest', 'MIT', 'Score distribution charts'],
        ['React Router', 'Latest', 'MIT', 'Page routing'],
        ['Ant Design / MUI', 'Latest', 'MIT', 'UI component library'],
    ]
)

para('Total Package Cost: ₹0 (all open-source, except PyMuPDF commercial license if needed)', bold=True)

doc.add_page_break()

# ==================== 9. HARDWARE ====================
heading('9. HARDWARE REQUIREMENTS (Development Machines)', 1)
table(
    ['#', 'Item', 'Minimum Spec', 'Recommended Spec', 'Purpose', 'Priority'],
    [
        ['1', 'Developer Laptop/PC', '8GB RAM, i5, 256GB SSD', '16GB RAM, i7/Ryzen 7, 512GB SSD', 'Development & local testing', 'P0'],
        ['2', 'Internet Connection', '10 Mbps', '50+ Mbps', 'Cloud access, model downloads', 'P0'],
        ['3', 'External Monitor (Optional)', '—', '24" 1080p', 'Productivity', 'P3'],
    ]
)
para('Note: Models (DistilBERT ~260MB, SBERT ~80MB) need to be downloaded once. Stable internet required.', bold=True)

doc.add_page_break()

# ==================== 10. TEAM ACCESS ====================
heading('10. TEAM ACCESS & STAKEHOLDERS', 1)

heading('10.1 Development Team Access', 2)
table(
    ['#', 'Access Required', 'For Whom', 'Purpose', 'Priority'],
    [
        ['1', 'AWS Console Access', 'All developers', 'Infrastructure management', 'P0'],
        ['2', 'GitHub Repo Access', 'All developers', 'Code push/pull', 'P0'],
        ['3', 'Database Read/Write', 'Backend developers', 'Schema design & queries', 'P0'],
        ['4', 'S3 Bucket Access', 'Backend developers', 'Resume upload/download', 'P0'],
        ['5', 'Redis Access', 'Backend developers', 'Queue management', 'P1'],
        ['6', 'Monitoring Dashboard', 'DevOps / Lead', 'System health', 'P2'],
    ]
)

heading('10.2 Stakeholder Access', 2)
table(
    ['#', 'Person/Role', 'What They Provide', 'When Needed', 'Priority'],
    [
        ['1', 'HR Manager (1–2 people)', 'Feedback on dashboard, weight calibration, UAT testing', 'Week 3 onwards', 'P0'],
        ['2', 'Placement Cell Coordinator', 'Student data, past resumes, eligibility rules', 'Day 1', 'P0'],
        ['3', 'Legal/Compliance Officer', 'Data privacy approval, PII handling guidelines', 'Before Day 1', 'P0'],
        ['4', 'IT/Infra Admin', 'Server provisioning, firewall rules, VPN', 'Day 1', 'P0'],
        ['5', 'UI/UX Designer (Optional)', 'Dashboard design mockups', 'Week 2', 'P3'],
        ['6', 'QA Tester', 'End-to-end testing, parsing validation', 'Week 5 onwards', 'P2'],
    ]
)

doc.add_page_break()

# ==================== 11. COMPLIANCE ====================
heading('11. COMPLIANCE & LEGAL APPROVALS', 1)
table(
    ['#', 'Approval Required', 'From Whom', 'Why', 'Priority'],
    [
        ['1', 'Data Processing Approval', 'Legal / DPO', 'Permission to process student resumes & PII', 'P0'],
        ['2', 'Data Retention Policy', 'Legal / Management', 'How long to store resumes & scores', 'P1'],
        ['3', 'Student Consent Framework', 'Legal', 'Students consent to AI-based resume analysis', 'P0'],
        ['4', 'GDPR / IT Act Compliance', 'Legal', 'Ensure compliance with data protection laws', 'P1'],
        ['5', 'Open-Source License Review', 'Legal', 'Verify PyMuPDF AGPL license for commercial use', 'P2'],
        ['6', 'AI Transparency Policy', 'Management', 'Disclose that AI scoring is used in shortlisting', 'P1'],
        ['7', 'Audit Trail Requirements', 'Compliance', 'Define what must be logged and for how long', 'P2'],
    ]
)

doc.add_page_break()

# ==================== 12. COST SUMMARY ====================
heading('12. COMPLETE COST SUMMARY', 1)

heading('12.1 One-Time Costs', 2)
table(
    ['#', 'Item', 'Cost (INR)', 'Notes'],
    [
        ['1', 'GPU for Model Training', '₹500 – ₹1,500', 'Google Colab Pro or 4-8 hrs EC2 GPU'],
        ['2', 'Domain Name (Annual)', '₹500 – ₹1,500', 'One-time purchase, annual renewal'],
        ['3', 'SSL Certificate (if commercial)', '₹0 – ₹5,000', 'Free with Let\'s Encrypt'],
        ['4', 'PyMuPDF Commercial License', '₹0 – ₹50,000', 'Only if AGPL not acceptable'],
        ['', 'Total One-Time', '₹1,000 – ₹58,000', 'Depends on license decisions'],
    ]
)

heading('12.2 Monthly Recurring Costs', 2)
table(
    ['#', 'Item', 'Monthly Cost (INR)', 'Priority'],
    [
        ['1', 'EC2 Compute (t3.large)', '₹6,000', 'P0'],
        ['2', 'PostgreSQL (RDS)', '₹2,000', 'P0'],
        ['3', 'Redis (ElastiCache)', '₹1,200', 'P1'],
        ['4', 'S3 Storage', '₹500', 'P0'],
        ['5', 'Monitoring (CloudWatch)', '₹500', 'P2'],
        ['6', 'Backups & DR', '₹500', 'P2'],
        ['7', 'SSL / Domain renewal', '₹200', 'P1'],
        ['', '', '', ''],
        ['', 'SUBTOTAL (Core)', '₹10,900', ''],
        ['', '', '', ''],
        ['8', 'OpenAI API (Optional)', '₹1,200', 'P3'],
        ['9', 'Email Service (SES)', '₹200', 'P2'],
        ['10', 'GitHub Team (if needed)', '₹350/user', 'P1'],
        ['', '', '', ''],
        ['', 'TOTAL (with all optionals)', '₹12,650 + team costs', ''],
    ]
)

heading('12.3 Cost Per Resume', 2)
table(
    ['Scenario', 'Monthly Cost', 'Resumes/Month', 'Cost Per Resume'],
    [
        ['Core only', '₹10,900', '3,000', '₹3.63'],
        ['With LLM', '₹12,100', '3,000', '₹4.03'],
        ['High volume', '₹10,900', '6,000', '₹1.82'],
        ['Low volume', '₹10,900', '1,000', '₹10.90'],
    ]
)

heading('12.4 Zero-Cost Items (Open Source)', 2)
table(
    ['Category', 'Items', 'Total Cost'],
    [
        ['ML Models', 'DistilBERT, Sentence-BERT, spaCy', '₹0'],
        ['Backend Framework', 'FastAPI, Celery, SQLAlchemy', '₹0'],
        ['Frontend Framework', 'React, Vite, Recharts', '₹0'],
        ['DevOps', 'Docker, Docker Compose, GitHub Actions (free tier)', '₹0'],
        ['OCR', 'EasyOCR, Tesseract', '₹0'],
        ['Monitoring', 'Prometheus, Grafana (self-hosted)', '₹0'],
        ['Annotation', 'Label Studio (self-hosted)', '₹0'],
        ['API Docs', 'Swagger (auto-generated)', '₹0'],
    ]
)

doc.add_page_break()

# ==================== 13. PRIORITY SUMMARY ====================
heading('13. PRIORITY-WISE SUMMARY FOR PROJECT MANAGER', 1)

heading('P0 — MUST HAVE BEFORE DAY 1', 2)
para('Without these, development CANNOT start.', bold=True)
rows_p0 = [
    ['1', '500–1,000 anonymized student resumes (PDF)', 'Placement Cell', '⭐ MOST CRITICAL'],
    ['2', '20–50 sample job descriptions', 'HR / Placement Cell', '⭐ CRITICAL'],
    ['3', 'Student profile data (branch, CGPA, batch)', 'Placement Cell', 'CRITICAL'],
    ['4', 'AWS / GCP cloud account with billing', 'IT Admin', 'CRITICAL'],
    ['5', 'EC2 t3.large instance', 'IT Admin / Cloud', 'CRITICAL'],
    ['6', 'S3 storage bucket', 'IT Admin / Cloud', 'CRITICAL'],
    ['7', 'PostgreSQL database', 'IT Admin / Cloud', 'CRITICAL'],
    ['8', 'GPU access (one-time, 4–8 hrs)', 'Cloud / Colab Pro', 'CRITICAL'],
    ['9', 'GitHub private repository', 'IT Admin', 'CRITICAL'],
    ['10', 'IAM credentials for developers', 'IT Admin', 'CRITICAL'],
    ['11', 'Data processing approval (legal)', 'Legal / DPO', 'CRITICAL'],
    ['12', 'Student consent framework', 'Legal', 'CRITICAL'],
    ['13', 'HR stakeholder (1–2 people)', 'HR Department', 'CRITICAL'],
    ['14', 'Placement Cell coordinator', 'Placement Cell', 'CRITICAL'],
    ['15', 'Developer laptops (16GB RAM)', 'IT Admin', 'CRITICAL'],
]
table(['#', 'Resource', 'Request From', 'Status'], rows_p0)

heading('P1 — NEED WITHIN WEEK 1', 2)
rows_p1 = [
    ['1', 'Redis instance (cache + queue)', 'Cloud / IT Admin'],
    ['2', 'Domain name + SSL certificate', 'IT Admin'],
    ['3', 'Docker Desktop + registry access', 'IT Admin'],
    ['4', 'Jira / Trello project board', 'PM / IT Admin'],
    ['5', 'Slack / Teams project channel', 'IT Admin'],
    ['6', 'Firewall rules configured', 'IT Admin'],
    ['7', 'Data retention policy', 'Legal'],
    ['8', 'AI transparency disclosure policy', 'Management'],
]
table(['#', 'Resource', 'Request From'], rows_p1)

heading('P2 — NEED WITHIN WEEK 2–3', 2)
rows_p2 = [
    ['1', 'CI/CD pipeline (GitHub Actions)', 'DevOps'],
    ['2', 'CloudWatch / Prometheus monitoring', 'DevOps / Cloud'],
    ['3', 'Email service (AWS SES / SendGrid)', 'IT Admin'],
    ['4', 'Automated database backups', 'DBA / Cloud'],
    ['5', 'ClamAV for malware scanning', 'IT Security'],
    ['6', 'QA tester assigned', 'PM'],
    ['7', 'Open-source license review', 'Legal'],
    ['8', 'Audit trail requirements defined', 'Compliance'],
]
table(['#', 'Resource', 'Request From'], rows_p2)

heading('P3 — NICE TO HAVE', 2)
rows_p3 = [
    ['1', 'OpenAI API key (for AI summaries)', 'PM / Finance'],
    ['2', 'Historical hiring data', 'HR'],
    ['3', 'Staging server (t3.medium)', 'Cloud'],
    ['4', 'UI/UX designer', 'PM'],
    ['5', 'CDN (CloudFront)', 'Cloud'],
    ['6', 'PyCharm Professional licenses', 'IT Admin'],
    ['7', 'Weights & Biases account', 'DevOps'],
]
table(['#', 'Resource', 'Request From'], rows_p3)

doc.add_paragraph()
note('ACTION ITEM: Start with requesting P0 items immediately. The 500+ sample resumes will take the longest to collect — start this process TODAY.', 'PM ACTION')

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Document Version: 1.0 | Date: 26 February 2026 | Project: CRMS Resume Intelligence System')
r.font.size = Pt(9)
r.font.color.rgb = RGBColor(128, 128, 128)

# ==================== SAVE ====================
output = r'c:\Dev\CRMS_Resume\CRMS_Resource_Checklist.docx'
doc.save(output)
print(f'Word doc saved: {output}')

from docx2pdf import convert
pdf_path = r'c:\Dev\CRMS_Resume\CRMS_Resource_Checklist.pdf'
convert(output, pdf_path)
print(f'PDF saved: {pdf_path}')
